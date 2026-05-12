from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os
from utils.setup_logger_utils import setup_logger
from docx.enum.section import WD_ORIENT
from PIL import Image

# ----------------------- Setup Logging -----------------------
logger = setup_logger("word_utils_logger", "word_utils.log")


# ---------------- DOCUMENT ----------------
def create_document() -> Document:
    try:
        document = Document()

        section = document.sections[0]

        # A4 size
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        # margins
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        return document

    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise

# ---------------- TITLE ----------------
def add_folder_title(document: Document, folder_name: str):
    try:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(folder_name)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 30, 30)
        run.font.name = "Arial"

    except Exception as e:
        logger.error(f"Error adding folder title '{folder_name}': {e}")
        raise


# ---------------- TABLE STYLE ----------------
def style_table(table):
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr

        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "2")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "C0C0C0")
            tblBorders.append(border)

        tblPr.append(tblBorders)

    except Exception as e:
        logger.error(f"Error styling table: {e}")
        raise


def create_placeholder(width_px=600, height_px=350):
    img = Image.new("RGB", (width_px, height_px), (255, 255, 255))
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


# ---------------- CELL ----------------
def render_image_cell(cell, title: str, image: BytesIO):
    try:
        # remove everything inside cell
        cell._element.clear_content()

        # create single paragraph
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # remove ALL spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        # remove cell margins
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        for tag in ("top", "bottom", "left", "right"):
            elem = OxmlElement(f"w:{tag}")
            elem.set(qn("w:w"), "0")
            elem.set(qn("w:type"), "dxa")
            tcPr.append(elem)

        # use placeholder if image missing
        if not image:
            image = create_placeholder()
            logger.warning(f"No image found for {title}")
        else:
            image.seek(0)

        run = p.add_run()
        run.add_picture(image, width=Inches(2.38))

    except Exception as e:
        logger.error(f"Error rendering cell '{title}': {e}")
        raise

# ---------------- STUDENT CARD ----------------
def add_student_table(document: Document, guardian_img: BytesIO, national_img: BytesIO):
    try:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        style_table(table)

        cells = table.rows[0].cells
        render_image_cell(cells[1], "صورة بطاقة ولي الأمر", guardian_img)
        render_image_cell(cells[0], "صورة بطاقة الطالب", national_img)

    except Exception as e:
        logger.error(f"Error adding student table: {e}")
        raise


# ---------------- SAVE ----------------
def save_document(document: Document, output_dir: str = "output") -> str:
    try:
        os.makedirs(output_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"processed_national_ids_{timestamp}.docx"
        full_path = os.path.join(output_dir, filename)

        document.save(full_path)

        return full_path

    except Exception as e:
        logger.error(f"Error saving document '{filename}': {e}")
        raise